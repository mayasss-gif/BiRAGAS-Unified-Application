#!/usr/bin/env python3
"""
Generate Professional PDF + DOCX Manual for BiRAGAS CRISPR Complete v3.0
=========================================================================
Ayass Bioscience LLC — Unified DNA + RNA CRISPR Analysis Platform
"""

import fitz  # PyMuPDF for PDF
import docx  # python-docx for DOCX
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, datetime

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
TODAY = datetime.date.today().strftime("%B %d, %Y")

# ══════════════════════════════════════════════════════════════════════════════
# CONTENT DATA — shared between PDF and DOCX
# ══════════════════════════════════════════════════════════════════════════════

TITLE = "BiRAGAS CRISPR Complete v3.0"
SUBTITLE = "User Manual & Technical Guide — DNA + RNA"
COMPANY = "Ayass Bioscience LLC"

TOC = [
    "1. Overview & Capabilities",
    "2. Installation & Setup",
    "3. Quick Start Guide",
    "4. System Architecture (130 Python Files)",
    "5. EditingEngine — DNA + RNA Guide Design",
    "6. ScreeningEngine — MAGeCK + BAGEL2 + DrugZ + Perturb-seq",
    "7. KnockoutEngine — 7-Method Ensemble (421,718 Configs)",
    "8. MegaScaleEngine — Sparse Matrix O(1)",
    "9. CombinationEngine — 12-Model Cross-Modal (88.9B)",
    "10. ACE Scoring — 15-Stream Evidence Aggregation",
    "11. RNA Engines — Cas13 + Base Editing + ncRNA",
    "12. CRISPRi/CRISPRa — Transcriptome Modulation",
    "13. 7-Phase Causality Framework (28 Modules)",
    "14. 23 Core Causality Modules",
    "15. Autonomous Agentic AI Systems",
    "16. REST API Reference (13 Endpoints)",
    "17. Web Application Guide (10 Tabs)",
    "18. RNA Target Types (9 Classes)",
    "19. Scoring Formulas & Algorithms",
    "20. Configuration Reference",
    "21. BiRAGAS CRISPR vs CRIS.py",
    "22. Glossary",
]

SCALE_TABLE = [
    ["DNA Knockout Configs", "210,859", "19,169 genes x 11 configs"],
    ["RNA Knockdown Configs", "210,859", "19,169 genes x 11 configs"],
    ["Total Configs", "421,718", "DNA + RNA combined"],
    ["DNA x DNA Combinations", "22.2 Billion", "Knockout x Knockout"],
    ["RNA x RNA Combinations", "22.2 Billion", "Knockdown x Knockdown"],
    ["DNA x RNA Cross-Modal", "44.5 Billion", "Knockout x Knockdown"],
    ["Total Combinations", "88.9 Billion", "All modality pairs"],
    ["Knockout Methods", "7", "Weighted ensemble"],
    ["Synergy Models", "12", "6 classical + 6 cross-modal"],
    ["ACE Streams", "15", "Multi-modal evidence"],
    ["Causality Phases", "7", "28 integration modules"],
    ["Core Causality Modules", "23", "Pearl's do-calculus"],
    ["RNA Target Types", "9", "mRNA to spatial RNA"],
    ["Cas13 Variants", "4", "Cas13a/b/d + dCas13"],
    ["Python Files", "130", "Complete codebase"],
    ["Autonomous Checks", "9", "Self-correction"],
]

DNA_NUCLEASES = [
    ["SpCas9", "NGG", "20 nt", "Double-strand break", "Standard knockout"],
    ["SaCas9", "NNGRRT", "21 nt", "Double-strand break", "Smaller Cas9 (AAV)"],
    ["Cas12a", "TTTV", "23 nt", "Staggered cut", "AT-rich regions"],
    ["Cas12a", "TTTN", "23 nt", "Staggered cut", "Broad targeting"],
]

RNA_NUCLEASES = [
    ["Cas13a (LwaCas13a)", "H (not G)", "28 nt", "Yes", "Diagnostics (SHERLOCK)"],
    ["Cas13b (PspCas13b)", "NAN", "30 nt", "Yes", "RNA knockdown"],
    ["Cas13d (CasRx)", "None", "22 nt", "No", "Therapeutic KD (best)"],
    ["dCas13 (dead)", "None", "22 nt", "No", "Base editing, imaging, splicing"],
]

RNA_TYPES = [
    ["mRNA", "Protein-coding messenger RNA", "Cas13d (CasRx)", "Perturb-seq, bulk/scRNA-seq", "210,859 configs"],
    ["lncRNA", "Long non-coding RNA (>200nt)", "CRISPRi at TSS", "scRNA-seq, CHART-seq", "Per-locus strategy"],
    ["miRNA", "MicroRNA (~22nt regulator)", "Cas13d / Cas9 pre-miRNA KO", "Small RNA-seq, CLIP-seq", "Seed derepression"],
    ["siRNA", "Small interfering RNA (~21nt)", "Cas13d direct targeting", "Small RNA-seq", "Direct degradation"],
    ["circRNA", "Circular RNA (backsplice)", "Cas13d / Cas9 backsplice KO", "RNase R + RNA-seq", "Junction targeting"],
    ["piRNA", "Piwi-interacting RNA (~26-31nt)", "Cas13d / CRISPRi cluster", "piRNA-seq", "Cluster silencing"],
    ["scRNA", "Single-cell transcriptome", "Perturb-seq / CROP-seq", "10x Chromium scRNA-seq", "Cell-level profiling"],
    ["Bulk RNA", "Population transcriptome", "CRISPRi/CRISPRa screen", "Bulk RNA-seq (DESeq2)", "Population average"],
    ["Spatial RNA", "Tissue-resolved RNA", "CRISPR-TO", "MERFISH, Visium", "Spatial localization"],
]

KO_METHODS = [
    ["1", "Topological", "20%", "Pearl's do-calculus graph surgery"],
    ["2", "Bayesian", "18%", "Noisy-OR network propagation"],
    ["3", "Monte Carlo", "18%", "N=1000 stochastic simulations + 95% CI"],
    ["4", "Pathway", "14%", "Pathway-specific decay rates"],
    ["5", "Feedback", "12%", "Compensatory loop adjustment"],
    ["6", "ODE", "10%", "Euler method differential dynamics"],
    ["7", "Mutual Info", "8%", "Information-theoretic scoring"],
]

SYNERGY_12 = [
    ["1", "Bliss Independence", "Classical", "C = A + B - AB", "20-22%"],
    ["2", "HSA", "Classical", "C = max(A, B)", "10-15%"],
    ["3", "Loewe Additivity", "Classical", "Isobole: dA/DA + dB/DB = 1", "15-20%"],
    ["4", "ZIP", "Classical", "AB + network proximity", "12-18%"],
    ["5", "Graph Epistasis", "Classical", "DAG descendant overlap", "12-15%"],
    ["6", "Compensation Blocking", "Classical", "Resistance pathway targeting", "8-10%"],
    ["7", "Transcriptional Cascade", "Cross-Modal", "KO removes TF + KD removes transcript", "5-12%"],
    ["8", "Isoform Escape Blocker", "Cross-Modal", "KO isoform A + KD isoform B", "5-8%"],
    ["9", "Feedback Loop Disruptor", "Cross-Modal", "KO upstream + KD feedback RNA", "5-7%"],
    ["10", "Collateral Synergy", "Cross-Modal", "Cas13a/b collateral + partner hit", "0-5%"],
    ["11", "Epigenetic-Transcriptomic", "Cross-Modal", "CRISPRi + Cas13 degradation", "0-5%"],
    ["12", "ncRNA-Coding Network", "Cross-Modal", "lncRNA/miRNA + coding gene", "5%"],
]

ACE_STREAMS = [
    ["1", "MAGeCK RRA", "0.90", "RRA p-value", "DNA Screen"],
    ["2", "MAGeCK MLE", "0.85", "MLE beta", "DNA Screen"],
    ["3", "BAGEL2 BF", "0.85", "Bayes Factor", "Essentiality"],
    ["4", "Perturb-seq", "0.80", "Effect size", "scRNA-seq"],
    ["5", "GWAS", "0.90", "p-value", "Genomic"],
    ["6", "MR beta", "0.95", "Mendelian Randomization", "Causal"],
    ["7", "eQTL", "0.85", "Association", "Expression"],
    ["8", "SIGNOR", "0.90", "Network weight", "Pathway"],
    ["9", "Centrality", "0.70", "Betweenness", "Network"],
    ["10", "DAG Tier", "0.75", "Network tier", "Topology"],
    ["11", "Drug Sensitivity", "0.80", "DrugZ", "Drug"],
    ["12", "Conservation", "0.65", "Phylogenetic", "Evolution"],
    ["13", "FluteMLE Pathway", "0.70", "Enrichment", "RNA Pathway"],
    ["14", "Editing Efficiency", "0.60", "KD/KO success", "RNA/DNA Edit"],
    ["15", "Drug Z-score", "0.75", "Z-score", "Drug"],
]

PHASE_MODULES = [
    ["Phase 1", "Screening to DAG", "QualityGate", "ScreeningConverter", "CRISPREnricher", "MRCorroborator"],
    ["Phase 2", "Network Scoring", "TargetScorer(7D)", "CentralityEnhancer", "TierPromoter", "KnockoutBridge"],
    ["Phase 3", "Quality Assurance", "DirectionValidator", "ConfoundingDetector", "HallucinationShield", "KOIntegrator"],
    ["Phase 4", "Mechanisms", "AttributeHarmonizer", "ResistanceEnhancer", "CompensationBridge", "EngineAdapter"],
    ["Phase 5", "Pharmaceutical", "DrugTargetRanker(9D)", "SafetyEnhancer", "EfficacyInjector", "SynergyUpgrader"],
    ["Phase 6", "Stratification", "WeightedStratifier", "DriverComparator", "MotifValidator", "SubtypeMapper"],
    ["Phase 7", "Reporting", "ReportGenerator", "QualityBooster", "ArbitrationEnhancer", "GapPrioritizer"],
]

CAUSALITY_23 = [
    ["Phase 1", "dag_builder", "1,191", "Build consensus DAG from 9 data streams"],
    ["Phase 1", "mr_validator", "325", "5 MR methods + 4 sensitivity analyses"],
    ["Phase 2", "centrality_calculator", "215", "5 metrics + 3-tier classification"],
    ["Phase 2", "target_scorer", "165", "5-dimension composite scoring"],
    ["Phase 3", "causality_tester", "292", "5 hallucination categories, dual gate"],
    ["Phase 3", "directionality_tester", "205", "ACE asymmetry direction testing"],
    ["Phase 3", "confounding_checker", "200", "Pearl's backdoor criterion"],
    ["Phase 4", "counterfactual_simulator", "181", "do-calculus graph surgery"],
    ["Phase 4", "dag_propagation_engine", "172", "Forward/backward propagation"],
    ["Phase 4", "resistance_mechanism_identifier", "240", "5 resistance types"],
    ["Phase 4", "compensation_pathway_analyzer", "222", "4 compensation types"],
    ["Phase 5", "druggability_scorer", "149", "12 protein families"],
    ["Phase 5", "efficacy_predictor", "153", "Counterfactual-based efficacy"],
    ["Phase 5", "safety_assessor", "196", "8 critical pathways"],
    ["Phase 5", "target_ranker", "162", "7-dimension ranking"],
    ["Phase 5", "combination_analyzer", "219", "Synergy/coverage/safety"],
    ["Phase 6", "cohort_stratifier", "173", "Jaccard + k-medoids"],
    ["Phase 6", "dag_comparator", "219", "Structural/driver/pathway comparison"],
    ["Phase 6", "conserved_motifs_identifier", "240", "5 motif types"],
    ["Phase 7", "evidence_inspector", "227", "Bayesian quality formula"],
    ["Phase 7", "gap_analyzer", "261", "5 gap categories"],
    ["Phase 7", "llm_arbitrator", "298", "Rule-based + LLM conflict resolution"],
    ["Phase 7", "response_formatter", "285", "5 output formats"],
]

COMPARISON = [
    ["Guide Design", "None", "Multi-PAM DNA (NGG/NNGRRT/TTTV) + Cas13a/b/d RNA"],
    ["RNA Target Types", "None", "9 types: mRNA, lncRNA, miRNA, siRNA, circRNA, piRNA, scRNA, bulk, spatial"],
    ["Knockout Configs", "Manual", "421,718 (210,859 DNA + 210,859 RNA)"],
    ["Combinations", "None", "88.9B (22.2B DNA + 22.2B RNA + 44.5B cross-modal)"],
    ["Synergy Models", "None", "12 (6 classical + 6 cross-modal DNA x RNA)"],
    ["ACE Scoring", "None", "15-stream with Bayesian confidence"],
    ["RNA Base Editing", "None", "A-to-I (ADAR2) + C-to-U (APOBEC)"],
    ["CRISPRi/CRISPRa", "None", "Transcriptome modulation (no DNA cut)"],
    ["Perturb-seq", "None", "Single-cell CRISPR screen analysis"],
    ["Non-coding RNA", "None", "lncRNA/miRNA/siRNA/circRNA/piRNA"],
    ["Causality", "None", "7-phase BiRAGAS (23 modules + 28 integration)"],
    ["Self-Correction", "None", "9 autonomous checks + auto-fix + retry"],
    ["Web UI", "None", "10-tab professional interface (DNA + RNA)"],
    ["License", "GPL", "Proprietary (Ayass Bioscience LLC)"],
]

GLOSSARY = [
    ("ACE", "Aggregated CRISPR Evidence - 15-stream composite score"),
    ("ADAR2", "Adenosine Deaminase Acting on RNA 2 - A-to-I editing enzyme"),
    ("APOBEC", "Apolipoprotein B mRNA Editing Catalytic - C-to-U editing"),
    ("BAGEL2", "Bayesian Analysis of Gene EssentiaLity"),
    ("Bliss Independence", "Combination model: C = A + B - AB"),
    ("Brunello", "Genome-scale sgRNA library (77,441 guides, 19,169 genes)"),
    ("Cas13", "RNA-targeting CRISPR nuclease family (a/b/d variants)"),
    ("circRNA", "Circular RNA formed by backsplicing"),
    ("CRISPRi", "CRISPR interference - transcriptional repression via dCas9-KRAB"),
    ("CRISPRa", "CRISPR activation - transcriptional activation via dCas9-VPR"),
    ("CRISPR-TO", "CRISPR perturbation + spatial transcriptomics"),
    ("CROP-seq", "CRISPR droplet sequencing (guide ID + transcriptome)"),
    ("DAG", "Directed Acyclic Graph - causal network structure"),
    ("dCas13", "Catalytically dead Cas13 - binds RNA without cutting"),
    ("do-calculus", "Pearl's mathematical framework for causal intervention"),
    ("DrugZ", "Drug screen Z-score analysis tool"),
    ("HDR", "Homology-Directed Repair - precise DNA editing"),
    ("lncRNA", "Long non-coding RNA (>200 nucleotides)"),
    ("MAGeCK", "Model-based Analysis of Genome-wide CRISPR-Cas9 Knockout"),
    ("miRNA", "MicroRNA (~22nt post-transcriptional regulator)"),
    ("PAM", "Protospacer Adjacent Motif - nuclease recognition site"),
    ("Perturb-seq", "CRISPR perturbation + single-cell RNA-seq readout"),
    ("PFS", "Protospacer Flanking Site - RNA equivalent of PAM (Cas13)"),
    ("piRNA", "Piwi-interacting RNA (~26-31nt, transposon defense)"),
    ("scCLEAN", "CRISPR-based rRNA depletion for scRNA-seq"),
    ("scRNA-seq", "Single-cell RNA sequencing"),
    ("sgRNA", "Single guide RNA - directs Cas9/Cas12a to DNA target"),
    ("siRNA", "Small interfering RNA (~21nt)"),
    ("ZIP", "Zero Interaction Potency - synergy reference model"),
]

# ══════════════════════════════════════════════════════════════════════════════
# DOCX GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def build_docx():
    doc = docx.Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E) if level == 1 else RGBColor(0x15, 0x65, 0xC0)
        return h

    def add_para(text, bold=False, italic=False, size=10, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_table(headers, rows, widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
        doc.add_paragraph()
        return t

    # ── Cover Page ──
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Unified DNA + RNA CRISPR Analysis Platform")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("User Manual & Technical Guide")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COMPANY)
    run.bold = True
    run.font.size = Pt(16)

    add_para(f"\nDocument Date: {TODAY}", size=10, color=RGBColor(0x5F, 0x63, 0x68))
    add_para("Classification: Proprietary & Confidential", size=10, color=RGBColor(0xD3, 0x2F, 0x2F))

    doc.add_page_break()

    # ── Table of Contents ──
    add_heading("Table of Contents", 1)
    for item in TOC:
        add_para(item, size=11)
    doc.add_page_break()

    # ── Ch 1: Overview ──
    add_heading("1. Overview & Capabilities", 1)
    add_para("BiRAGAS CRISPR Complete v3.0 is the first platform to unify DNA editing (Cas9/Cas12a) and RNA targeting (Cas13/dCas13/CRISPRi/CRISPRa) with 7-phase causal inference for drug discovery.", size=10)
    add_para("\nScale & Capabilities:", bold=True, size=11)
    add_table(["Metric", "Value", "Detail"], SCALE_TABLE)

    # ── Ch 2: Installation ──
    add_heading("2. Installation & Setup", 1)
    add_para("Requirements: Python 3.10+ (3.12 recommended), 8GB RAM minimum, numpy/scipy/networkx/fastapi/uvicorn", size=10)
    add_para("\npip install numpy scipy networkx fastapi uvicorn pydantic", bold=True, size=9)
    add_para("python -m BiRAGAS_CRISPR_Complete --serve  # Start web server at http://localhost:8000", size=9)

    # ── Ch 3: Quick Start ──
    add_heading("3. Quick Start Guide", 1)
    add_para("Option A: Double-click BiRAGAS_CRISPR_Complete_App.html (works offline)", bold=True, size=10)
    add_para("Option B: python -m BiRAGAS_CRISPR_Complete --serve (full backend)", bold=True, size=10)
    add_para("Option C: Python API:", bold=True, size=10)
    add_para('from BiRAGAS_CRISPR_Complete import UnifiedOrchestrator\norch = UnifiedOrchestrator()\nreport = orch.run(crispr_dir="/path/", disease_name="Melanoma")', size=9)

    # ── Ch 5: EditingEngine ──
    doc.add_page_break()
    add_heading("5. EditingEngine — DNA + RNA Guide Design", 1)
    add_para("Unified engine supporting both DNA nucleases and RNA nucleases in a single API.", size=10)
    add_para("\nDNA Nucleases:", bold=True, size=10)
    add_table(["Nuclease", "PAM", "Guide Length", "Cut Type", "Best For"], DNA_NUCLEASES)
    add_para("RNA Nucleases (Cas13 Family):", bold=True, size=10)
    add_table(["Variant", "PFS", "Guide Length", "Collateral", "Best For"], RNA_NUCLEASES)

    # ── Ch 7: KnockoutEngine ──
    add_heading("7. KnockoutEngine — 7-Method Ensemble", 1)
    add_para("Scale: 210,859 DNA knockout configs + 210,859 RNA knockdown configs = 421,718 total.", size=10)
    add_para("\n7 Propagation Methods:", bold=True, size=10)
    add_table(["#", "Method", "Weight", "Algorithm"], KO_METHODS)

    # ── Ch 9: CombinationEngine ──
    doc.add_page_break()
    add_heading("9. CombinationEngine — 12-Model Cross-Modal (88.9B)", 1)
    add_para("The first combination engine modeling DNA x DNA, RNA x RNA, and DNA x RNA cross-modal synergy.", size=10)
    add_para("\n12 Synergy Models:", bold=True, size=10)
    add_table(["#", "Model", "Type", "Formula/Mechanism", "Weight"], SYNERGY_12)

    # ── Ch 10: ACE ──
    add_heading("10. ACE Scoring — 15-Stream Evidence Aggregation", 1)
    add_table(["#", "Stream", "Weight", "Source", "Type"], ACE_STREAMS)
    add_para("Formula: ACE = Sigma(normalized_i x weight_i) / Sigma(weight_i)", size=9, bold=True)

    # ── Ch 11: RNA Engines ──
    doc.add_page_break()
    add_heading("11. RNA Engines", 1)
    add_heading("RNA Target Types (9 Classes)", 2)
    add_table(["RNA Type", "Description", "Recommended Tool", "Analysis", "Scale"], RNA_TYPES)

    add_heading("RNA Base Editing (dCas13)", 2)
    add_para("A-to-I (ADAR2dd): Edit window positions 18-30. Prefers UAG > AAG > CAG context.", size=10)
    add_para("C-to-U (APOBEC1): Edit window positions 15-25. Creates stop codons or recodes amino acids.", size=10)

    # ── Ch 13: Causality ──
    doc.add_page_break()
    add_heading("13. 7-Phase Causality Framework (28 Modules)", 1)
    add_table(["Phase", "Focus", "Module 1", "Module 2", "Module 3", "Module 4"], PHASE_MODULES)

    add_heading("RNA Evidence Through 7 Phases", 2)
    phases_rna = [
        ["Phase 1", "RNA knockdown + Perturb-seq data enriches DAG nodes"],
        ["Phase 2", "RNA evidence boosts centrality and tier promotion"],
        ["Phase 3", "Cas13 KD validates causal edges (hallucination shield)"],
        ["Phase 4", "RNA compensation + isoform switching resistance detection"],
        ["Phase 5", "RNA knockdown efficacy + cross-modal combination synergy"],
        ["Phase 6", "scRNA-seq for patient stratification + subtype-specific RNA targets"],
        ["Phase 7", "Dual DNA+RNA validated targets in clinical report"],
    ]
    add_table(["Phase", "RNA Integration"], phases_rna)

    # ── Ch 14: 23 Modules ──
    add_heading("14. 23 Core Causality Modules", 1)
    add_table(["Phase", "Module", "Lines", "Description"], CAUSALITY_23)

    # ── Ch 15: Autonomous ──
    doc.add_page_break()
    add_heading("15. Autonomous Agentic AI Systems", 1)
    agentic = [
        ["SelfCorrector", "9 DAG checks: cycle removal, orphan connection, confidence pruning, attribute injection, layer enforcement, weight normalization, duplicate removal, component merging, confidence harmonization"],
        ["PipelineDebugger", "Error classification + exponential backoff retry (2/4/8/16/30s) + fallback methods + state checkpointing"],
        ["LearningEngine", "Performance tracking + adaptive phase optimization + pattern recognition"],
        ["StressTestAgent", "17 differential diagnosis scenarios for clinical validation"],
        ["EngineSelector", "Automatic selection: Classical (<100 genes), Agentic (100-5K), MegaScale (>5K)"],
    ]
    add_table(["System", "Description"], agentic)

    # ── Ch 18: RNA Types ──
    add_heading("18. RNA Target Types (9 Classes)", 1)
    for rt in RNA_TYPES:
        add_para(f"{rt[0]}: {rt[1]}", bold=True, size=10)
        add_para(f"  Tool: {rt[2]} | Analysis: {rt[3]} | Scale: {rt[4]}", size=9)

    # ── Ch 21: Comparison ──
    doc.add_page_break()
    add_heading("21. BiRAGAS CRISPR Complete vs CRIS.py", 1)
    add_table(["Feature", "CRIS.py", "BiRAGAS CRISPR Complete v3.0"], COMPARISON)

    # ── Ch 22: Glossary ──
    add_heading("22. Glossary", 1)
    add_table(["Term", "Definition"], [[g[0], g[1]] for g in GLOSSARY])

    # ── Final ──
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{TITLE}\n{COMPANY}\nProprietary & Confidential\nCopyright (c) {datetime.date.today().year}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    path = os.path.join(OUT_DIR, "BiRAGAS_CRISPR_Complete_Manual_v3.0.docx")
    doc.save(path)
    print(f"DOCX saved: {path}")
    return path


# ══════════════════════════════════════════════════════════════════════════════
# PDF GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf():
    doc = fitz.open()
    NAVY = (0.10, 0.14, 0.49)
    BLUE = (0.08, 0.40, 0.75)
    WHITE = (1, 1, 1)
    BLACK = (0, 0, 0)
    GRAY = (0.37, 0.39, 0.41)
    LGRAY = (0.96, 0.97, 0.98)
    MGRAY = (0.85, 0.86, 0.87)
    PURPLE = (0.42, 0.10, 0.60)
    W, H = 612, 792
    ML, MR, MT, MB = 54, 54, 72, 60
    CW = W - ML - MR

    page_num = [0]
    y = [MT]

    def new_page():
        p = doc.new_page(width=W, height=H)
        page_num[0] += 1
        y[0] = MT
        if page_num[0] > 1:
            p.draw_line(fitz.Point(ML, 50), fitz.Point(W - MR, 50), color=BLUE, width=0.5)
            p.insert_text(fitz.Point(ML, 45), "BiRAGAS CRISPR Complete v3.0 — DNA + RNA", fontname="heit", fontsize=7.5, color=GRAY)
            p.insert_text(fitz.Point(W - MR - 80, 45), "Ayass Bioscience LLC", fontname="heit", fontsize=7.5, color=GRAY)
        p.draw_line(fitz.Point(ML, H - 40), fitz.Point(W - MR, H - 40), color=MGRAY, width=0.5)
        p.insert_text(fitz.Point(ML, H - 28), "Proprietary & Confidential", fontname="heit", fontsize=7, color=GRAY)
        p.insert_text(fitz.Point(W - MR - 30, H - 28), f"Page {page_num[0]}", fontname="helv", fontsize=7, color=GRAY)
        return p

    def check(needed):
        if y[0] + needed > H - MB:
            return new_page()
        return doc[-1]

    def heading(text, level=1):
        p = check(40)
        y[0] += 12 if level == 1 else 8
        if level == 1:
            p.draw_rect(fitz.Rect(ML, y[0] - 4, ML + 4, y[0] + 16), color=BLUE, fill=BLUE)
            p.insert_text(fitz.Point(ML + 10, y[0] + 12), text, fontname="hebo", fontsize=14, color=NAVY)
            y[0] += 24
            p.draw_line(fitz.Point(ML, y[0] - 4), fitz.Point(W - MR, y[0] - 4), color=MGRAY, width=0.5)
            y[0] += 6
        else:
            p.insert_text(fitz.Point(ML, y[0] + 10), text, fontname="hebo", fontsize=11, color=BLUE)
            y[0] += 18

    def para(text, fs=9.5, color=BLACK, bold=False):
        font = "hebo" if bold else "helv"
        cpl = int(CW / (fs * 0.52))
        lines = []
        for raw in text.split('\n'):
            while len(raw) > cpl:
                sp = raw[:cpl].rfind(' ')
                if sp <= 0: sp = cpl
                lines.append(raw[:sp])
                raw = raw[sp:].lstrip()
            lines.append(raw)
        for line in lines:
            p = check(fs + 3)
            p.insert_text(fitz.Point(ML, y[0] + fs), line[:100], fontname=font, fontsize=fs, color=color)
            y[0] += fs + 3

    def table(headers, rows, col_widths=None):
        nc = len(headers)
        if not col_widths:
            col_widths = [CW / nc] * nc
        rh = 15

        p = check(rh + 4)
        x = ML
        p.draw_rect(fitz.Rect(ML, y[0], ML + sum(col_widths), y[0] + rh), color=NAVY, fill=NAVY)
        for i, h in enumerate(headers):
            p.insert_text(fitz.Point(x + 3, y[0] + 10), str(h)[:25], fontname="hebo", fontsize=7, color=WHITE)
            x += col_widths[i]
        y[0] += rh

        for ri, row in enumerate(rows):
            p = check(rh + 2)
            x = ML
            bg = LGRAY if ri % 2 == 0 else WHITE
            p.draw_rect(fitz.Rect(ML, y[0], ML + sum(col_widths), y[0] + rh), color=MGRAY, fill=bg)
            for i, cell in enumerate(row):
                p.insert_text(fitz.Point(x + 3, y[0] + 10), str(cell)[:35], fontname="helv", fontsize=7, color=BLACK)
                x += col_widths[i]
            y[0] += rh
        y[0] += 6

    # ── Cover Page ──
    p = new_page()
    p.draw_rect(fitz.Rect(0, 0, W, 280), color=NAVY, fill=NAVY)
    p.draw_rect(fitz.Rect(0, 280, W, 286), color=BLUE, fill=BLUE)
    p.insert_text(fitz.Point(ML + 10, 100), "BiRAGAS CRISPR", fontname="hebo", fontsize=36, color=WHITE)
    p.insert_text(fitz.Point(ML + 10, 140), "Complete", fontname="hebo", fontsize=36, color=(0.39, 0.71, 0.96))
    p.insert_text(fitz.Point(ML + 10, 180), "DNA + RNA Analysis Platform", fontname="helv", fontsize=18, color=(0.80, 0.85, 0.95))
    p.insert_text(fitz.Point(ML + 10, 215), "User Manual & Technical Guide  |  Version 3.0", fontname="helv", fontsize=12, color=(0.70, 0.75, 0.85))

    stats = [("421,718", "Configs"), ("88.9B", "Combinations"), ("12", "Synergy Models"), ("9", "RNA Types")]
    for i, (v, l) in enumerate(stats):
        x = ML + 10 + i * 130
        p.insert_text(fitz.Point(x, 320), v, fontname="hebo", fontsize=16, color=BLUE)
        p.insert_text(fitz.Point(x, 336), l, fontname="helv", fontsize=8, color=GRAY)

    p.insert_text(fitz.Point(ML + 10, 420), COMPANY, fontname="hebo", fontsize=18, color=NAVY)
    p.insert_text(fitz.Point(ML + 10, 445), f"Document Date: {TODAY}", fontname="helv", fontsize=10, color=GRAY)
    p.insert_text(fitz.Point(ML + 10, 462), "Classification: Proprietary & Confidential", fontname="helv", fontsize=10, color=(0.83, 0.18, 0.18))
    p.draw_rect(fitz.Rect(0, H - 25, W, H), color=NAVY, fill=NAVY)

    # ── TOC ──
    new_page()
    heading("Table of Contents")
    for item in TOC:
        para(item, fs=10, color=NAVY)

    # ── Ch 1 ──
    new_page()
    heading("1. Overview & Capabilities")
    para("BiRAGAS CRISPR Complete v3.0 is the first platform to unify DNA editing and RNA targeting with 7-phase causal inference for drug discovery.")
    para("Scale & Capabilities:", bold=True)
    table(["Metric", "Value", "Detail"], SCALE_TABLE, [180, 100, 224])

    # ── Ch 5: Editing ──
    new_page()
    heading("5. EditingEngine — DNA + RNA Guide Design")
    para("DNA Nucleases:", bold=True)
    table(["Nuclease", "PAM", "Guide", "Cut Type", "Best For"], DNA_NUCLEASES, [90, 70, 50, 120, 174])
    para("RNA Nucleases (Cas13):", bold=True)
    table(["Variant", "PFS", "Guide", "Collateral", "Best For"], RNA_NUCLEASES, [120, 60, 50, 50, 224])

    # ── Ch 7: Knockout ──
    new_page()
    heading("7. KnockoutEngine — 7-Method Ensemble (421,718 Configs)")
    para("210,859 DNA knockout + 210,859 RNA knockdown = 421,718 total configurations.", bold=True)
    table(["#", "Method", "Weight", "Algorithm"], KO_METHODS, [25, 100, 50, 329])

    # ── Ch 9: Combinations ──
    new_page()
    heading("9. CombinationEngine — 12-Model Cross-Modal (88.9B)")
    para("88.9 BILLION: 22.2B DNA x DNA + 22.2B RNA x RNA + 44.5B DNA x RNA", bold=True)
    table(["#", "Model", "Type", "Mechanism", "Weight"], SYNERGY_12, [25, 140, 70, 190, 79])

    # ── Ch 10: ACE ──
    new_page()
    heading("10. ACE Scoring — 15-Stream Evidence Aggregation")
    table(["#", "Stream", "Weight", "Source", "Type"], ACE_STREAMS, [25, 110, 50, 150, 169])

    # ── Ch 11: RNA Types ──
    new_page()
    heading("11. RNA Target Types (9 Classes)")
    table(["Type", "Description", "Tool", "Analysis", "Scale"], RNA_TYPES, [55, 140, 120, 110, 79])

    # ── Ch 13: Phases ──
    new_page()
    heading("13. 7-Phase Causality Framework (28 Modules)")
    table(["Phase", "Focus", "Mod 1", "Mod 2", "Mod 3", "Mod 4"], PHASE_MODULES, [55, 90, 95, 95, 95, 74])

    # ── Ch 14: 23 Modules ──
    new_page()
    heading("14. 23 Core Causality Modules")
    table(["Phase", "Module", "Lines", "Description"], CAUSALITY_23, [55, 170, 40, 239])

    # ── Ch 21: Comparison ──
    new_page()
    heading("21. BiRAGAS CRISPR Complete vs CRIS.py")
    table(["Feature", "CRIS.py", "BiRAGAS CRISPR v3.0"], COMPARISON, [120, 60, 324])

    # ── Ch 22: Glossary ──
    new_page()
    heading("22. Glossary")
    table(["Term", "Definition"], [[g[0], g[1]] for g in GLOSSARY], [120, 384])

    # ── Back Cover ──
    p = new_page()
    p.draw_rect(fitz.Rect(0, 300, W, 480), color=NAVY, fill=NAVY)
    p.insert_text(fitz.Point(ML + 30, 360), "BiRAGAS CRISPR Complete v3.0", fontname="hebo", fontsize=20, color=WHITE)
    p.insert_text(fitz.Point(ML + 30, 390), "Unified DNA + RNA Analysis Platform", fontname="helv", fontsize=14, color=(0.70, 0.75, 0.90))
    p.insert_text(fitz.Point(ML + 30, 420), COMPANY, fontname="helv", fontsize=12, color=(0.80, 0.83, 0.90))
    p.insert_text(fitz.Point(ML + 30, 445), f"Copyright (c) {datetime.date.today().year}. All Rights Reserved.", fontname="helv", fontsize=9, color=(0.80, 0.83, 0.90))

    path = os.path.join(OUT_DIR, "BiRAGAS_CRISPR_Complete_Manual_v3.0.pdf")
    doc.save(path)
    doc.close()
    print(f"PDF saved: {path} ({page_num[0]} pages)")
    return path


if __name__ == '__main__':
    print("Generating manuals...")
    docx_path = build_docx()
    pdf_path = build_pdf()
    print(f"\nDone!")
    print(f"  PDF:  {pdf_path}")
    print(f"  DOCX: {docx_path}")
